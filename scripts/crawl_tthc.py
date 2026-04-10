#!/usr/bin/env python3
"""
TTHC Master Pipeline Tool (All-in-one Crawler) - Group theo Mapped Folders x CSV

Luồng Tự động hóa Dữ liệu HOÀN TOÀN MỚI:
1. Fetct List các Bộ/Ban/Ngành.
2. Paging qua API dvcquocgia từ từng cơ quan để móc tất cả các ID.
3. Xuất kết quả toàn bộ Data vào Folder tương ứng (VD: TTHC_IDs/BoCongAn/id_tthc.csv).
4. Tải file Word TTHC nội dung và chuyển thô sang Text dơ.
5. Gom file txt vào các folder Cơ quan tương ứng.
"""

import asyncio
import json
import argparse
import io
import re
import unicodedata
from pathlib import Path
from urllib.parse import urlencode

import aiohttp
import requests
import pandas as pd
from tqdm import tqdm
from tqdm.asyncio import tqdm as async_tqdm
from docx import Document

# ================= CẤU HÌNH ĐƯỜNG DẪN =================
IDS_DIR = Path("data/thutuchanhchinh/TTHC_IDs")
TXT_OUTPUT_DIR = Path("data/thutuchanhchinh/txt")
TXT_PROGRESS_FILE = Path("data/thutuchanhchinh/master_download_progress.txt")

# URL API
REST_API_URL = "https://thutuc.dichvucong.gov.vn/jsp/rest.jsp"
DOWNLOAD_URL = "https://thutuc.dichvucong.gov.vn/jsp/tthc/export/export_word_detail_tthc.jsp"

HEADERS = {
    "User-Agent": "Mozilla/5.0",
    "Content-Type": "application/x-www-form-urlencoded; charset=UTF-8",
    "Accept": "application/json, text/javascript, */*; q=0.01",
}

# ================= HELPER CHUẨN HOÁ TÊN =================
def format_agency_name(raw_name: str) -> str:
    if not raw_name: return "Unknown_Agency"
    nfkd_form = unicodedata.normalize('NFKD', raw_name)
    no_accent = u"".join([c for c in nfkd_form if not unicodedata.combining(c)])
    no_accent = no_accent.replace('đ', 'd').replace('Đ', 'D')
    words = re.sub(r'[^a-zA-Z0-9\s]', '', no_accent).split()
    return "".join(word.capitalize() for word in words)


# ================= BƯỚC 1: LẤY INTERNAL IDs GROUP THEO CƠ QUAN =================
async def fetch_agency_list(session, agency_type: int) -> list:
    params = {
        "service": "procedure_get_list_agency_by_type_service_v2",
        "provider": "dvcquocgia",
        "type": "ref",
        "loaicoquan": str(agency_type)
    }
    data = urlencode({"params": json.dumps(params)})
    try:
        async with session.post(REST_API_URL, data=data, headers=HEADERS, timeout=15) as response:
            if response.status == 200:
                result = await response.json()
                if isinstance(result, list):
                    return result
    except: pass
    return []

def build_api_payload(page_idx: int, agency_id: str, agency_type: str) -> str:
    params = {
        "service": "procedure_advanced_search_service_v2",
        "provider": "dvcquocgia",
        "type": "ref", "recordPerPage": 500, "pageIndex": page_idx,
        "is_connected": 0, "keyword": "", 
        "agency_type": agency_type,
        "impl_agency_id": str(agency_id),
        "object_id": "-1", "field_id": "-1", "impl_level_id": "-1",
    }
    return urlencode({"params": json.dumps(params)})

async def step1_fetch_ids():
    print("\n[STEP 1] ĐANG QUÉT API VÀ TẠO PHÂN LỚP THƯ MỤC CƠ QUAN...")
    IDS_DIR.mkdir(parents=True, exist_ok=True)
    
    async with aiohttp.ClientSession(headers=HEADERS) as session:
        # Lấy danh sách Bộ (Cấp Trung ương: type 0)
        agencies = await fetch_agency_list(session, 0)
        for ag in agencies: ag["a_type"] = "0"
            
        print(f"🎯 Đã tìm thấy {len(agencies)} Bộ Ban Ngành. Bắt đầu quét cạn dữ liệu...")
        
        # Thêm ProgressBar đẹp mắt cho phần Fetch
        pbar = tqdm(agencies, desc="Quét ID TTHC", bar_format="{l_bar}{bar} | {n_fmt}/{total_fmt} [{elapsed}<{remaining}, {rate_fmt}]")
        
        for agency in pbar:
            ag_id = agency.get("ID")
            ag_name_raw = agency.get("NAME")
            ag_type_str = agency.get("a_type")
            
            if not ag_id: continue
            
            # Khởi tạo thư mục và xác định file CSV riêng biệt
            ag_folder = format_agency_name(ag_name_raw)
            agency_output_dir = IDS_DIR / ag_folder
            agency_output_dir.mkdir(parents=True, exist_ok=True)
            agency_csv_file = agency_output_dir / "id_tthc.csv"
            
            # Tính năng Cập nhật Status Bar Tên Cơ Quan Đang Chạy
            pbar.set_postfix_str(f"Đang xử lý: {ag_folder[:15]}...")
            
            saved_codes = set()
            if agency_csv_file.exists():
                try:
                    df_cache = pd.read_csv(agency_csv_file)
                    if "PROCEDURE_CODE" in df_cache.columns:
                        saved_codes = set(df_cache["PROCEDURE_CODE"].astype(str).tolist())
                except: pass
            
            agency_new_records = []
            page = 1
            
            while True:
                payload = build_api_payload(page, ag_id, ag_type_str)
                try:
                    resp = requests.post(REST_API_URL, data=payload, headers=HEADERS, timeout=15)
                    if resp.status_code == 200:
                        data = resp.json()
                        if not isinstance(data, list) or len(data) == 0: break
                            
                        for item in data:
                            code = str(item.get("PROCEDURE_CODE"))
                            if code and code not in saved_codes:
                                item["AGENCY_FOLDER"] = ag_folder
                                item["AGENCY_RAW_NAME"] = ag_name_raw
                                agency_new_records.append(item)
                                saved_codes.add(code)
                        
                        if len(data) < 500: break
                        page += 1
                except Exception as e:
                    break
                    
            if agency_new_records:
                df_new = pd.DataFrame(agency_new_records)
                write_header = not agency_csv_file.exists()
                df_new.to_csv(agency_csv_file, mode='a', index=False, header=write_header, encoding="utf-8-sig")
                # Thể hiện số lượng mới load để User thấy phê hơn
                pbar.set_postfix_str(f"Đã lưu +{len(agency_new_records)} thủ tục ({ag_folder[:10]})")


# ================= BƯỚC 2: TẢI & CONVERT DOCS VÀO ĐÚNG FOLDER =================
def extract_text_from_docx(content: bytes) -> str:
    try:
        doc = Document(io.BytesIO(content))
        lines = [p.text for p in doc.paragraphs]
        for tbl in doc.tables:
            for row in tbl.rows:
                lines.extend(cell.text for cell in row.cells)
        return "\n".join(l for l in lines if l.strip())
    except:
        return ""

async def download_doc_txt(session, semaphore, code, t_id, ag_folder):
    async with semaphore:
        await asyncio.sleep(0.01)
        url = f"{DOWNLOAD_URL}?maTTHC={code}&idTTHC={t_id}"
        
        agency_dir = TXT_OUTPUT_DIR / ag_folder
        agency_dir.mkdir(parents=True, exist_ok=True)
        out_file = agency_dir / f"{code}.txt"
        
        if out_file.exists():
            return (code, True)
            
        try:
            async with session.get(url, headers={"User-Agent": "Mozilla/5.0"}, timeout=60) as resp:
                if resp.status == 200:
                    content = await resp.read()
                    if len(content) > 100:
                        text = extract_text_from_docx(content)
                        if text:
                            out_file.write_text(text, encoding="utf-8")
                            return (code, True)
        except: pass
    return (code, False)

async def step2_download_docs():
    print("\n[STEP 2] ĐANG TẢI & PHÂN BỔ DOCS VÀO THƯ MỤC CÁC CƠ QUAN...")
    TXT_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    if not IDS_DIR.exists():
        return print(f"Lỗi: Hãy chạy Step 1 trước (Không có {IDS_DIR})")
    
    mapping = []
    csv_files = list(IDS_DIR.glob("**/*.csv"))
    for csv_file in csv_files:
        try:
            df = pd.read_csv(csv_file)
            for _, row in df.iterrows():
                code = str(row.get("PROCEDURE_CODE", ""))
                t_id = str(row.get("ID", ""))
                folder = str(row.get("AGENCY_FOLDER", "ThuTucChung"))
                if code and t_id and code != "nan":
                    mapping.append((code, t_id, folder))
        except: pass
            
    completed = set(TXT_PROGRESS_FILE.read_text().split("\n")) if TXT_PROGRESS_FILE.exists() else set()
    remaining = [(c, idx, fldr) for c, idx, fldr in mapping if c not in completed]
    
    if not remaining:
        return print("Tất cả ID đã được tải file Doc an toàn.")
        
    semaphore = asyncio.Semaphore(30)
    async with aiohttp.ClientSession() as session:
        tasks = [download_doc_txt(session, semaphore, c, idx, fldr) for c, idx, fldr in remaining]
        
        with open(TXT_PROGRESS_FILE, "a") as f_prog:
            for result in async_tqdm(asyncio.as_completed(tasks), total=len(tasks), desc="Tải File Docs"):
                code, success = await result
                if success: f_prog.write(f"{code}\n")


# ================= THỰC THI CHÍNH =================
def main():
    parser = argparse.ArgumentParser(description="Tool tự động hóa luồng DVC TTHC (Nested CSV Mode)")
    parser.add_argument("--step", type=int, choices=[1, 2], help="Chạy một bước cụ thể")
    parser.add_argument("--all", action="store_true", help="Chạy tự động siêu tốc từ A-Z")
    args = parser.parse_args()

    if args.all or args.step == 1:
        asyncio.run(step1_fetch_ids())
    
    if args.all or args.step == 2:
        asyncio.run(step2_download_docs())

if __name__ == "__main__":
    main()
